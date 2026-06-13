import os
import pypdf
import docx

def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text from a PDF file using pypdf."""
    text = ""
    try:
        reader = pypdf.PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text

def extract_text_from_txt(file_path: str) -> str:
    """Extracts text from a plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"Error reading TXT {file_path}: {e}")
        return ""

def parse_resume(file_path: str) -> str:
    """Detects file type and extracts text accordingly."""
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
